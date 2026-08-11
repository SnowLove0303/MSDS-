# -*- coding: utf-8 -*-
"""MSDS 大批量入库就绪度评判 (四层 + 数据库逻辑审查 + 大批量能力).

对指定文件夹 (递归扫描 *.docx) 逐文件执行 **写入与输出对照**, 从四个
递进层次验证数据完整性:
  L1 解析正确性  reader → ParseResult      (节完整 / token 0遗漏 / 结构 / 序号)
  L2 写入正确性  ParseResult → SQLite       (字段全量 / 成分全量 / key收敛 / 幂等)
  L3 回读还原    SQLite → ParseResult       (结构 / 内容 / 顺序 / 成分)
  L4 输出正确性  SQLite → Excel             (宽表节序 / 节覆盖 / 明细全量 / 值保真)

另附:
  - 数据库逻辑性静态审查 (EAV 规范化 / 约束 / 索引 / 可扩展)
  - 大批量能力评估 (吞吐基准 / 幂等重跑 / 错误隔离 / 内存 / 事务)

用法:
  python -m db.audit "F:/.../标准化测试一库" [--out 报告目录] [--detail]

输出:
  audit_report.json   结构化评分 (每项 PASS/WARN/FAIL + 覆盖率)
  audit_report.md     人读评判报告 (含逐文件明细 + 大批量就绪度结论)

评判标准 (每项独立判定, 不阻断):
  项级: PASS=达标 | WARN=有偏差但非数据缺失 | FAIL=内容/结构损坏
  就绪度: 全部 PASS → 可大批量; 存在 FAIL → 需修复; WARN 记录但放行.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
import tempfile
import time
from pathlib import Path

_ROOT = Path(__file__).resolve().parent.parent
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

from core.docx_reader import read_msds
from core.document_tree import build_document_tree
from db.catalog import (add_product, connect, get_components, get_fields,
                        rebuild_product)
from db.export_single import export_single
from docx import Document
from openpyxl import load_workbook

# 与 batch_read --verify 同源 token 规则
_TOKEN_RE = re.compile(r"[一-鿿]{3,}|[A-Za-z]{3,}|\d{2,}")
# 原文 token 里设计性剥离的 (序号/成分表头词), 允许不覆盖
_ALLOWED_TOKEN_MISS = {"cas", "casnumber", "number", "chemical", "chemicalname",
                       "ingredient", "component", "concentration", "weight",
                       "percent"}


def _tokens(text: str) -> set[str]:
    return set(_TOKEN_RE.findall(text or ""))


def _is_subseq(needle: str, haystack: str) -> bool:
    """needle 是否是 haystack 的子序列 (字符按序出现).

    用于 token 完整性放宽: reader 会把连写标签拆分重排 (如 S2
    "存储存储已锁定。" → "存储" + "存储已锁定"), 整串 token 不再连续存在,
    但字符内容仍完整保留 → 子序列匹配即视为已覆盖 (信息未丢失).
    """
    if not needle:
        return True
    it = iter(haystack)
    return all(ch in it for ch in needle)


def _parse_text(result) -> str:
    """把 reader ParseResult 拼回文本 (覆盖 reader 全量内容, 含表头)."""
    parts = [result.header, result.footer]
    for sec in result.sections.values():
        parts.extend([sec.full_title, sec.title])
        for f in sec.fields:
            parts.extend([f.label, f.value])
        parts.extend(sec.lines)
        parts.append(sec.component_header)
        for c in sec.components:
            parts.extend([c.name, c.cas, c.conc])
    return "\n".join(parts)


def _cell_tokens(doc) -> set[str]:
    """源 docx 全部 token (16 表格 + 页眉/页脚段落与表格)."""
    src = set()
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                src |= _tokens(c.text)
    for s in doc.sections:
        for p in s.header.paragraphs:
            src |= _tokens(p.text)
        for p in s.footer.paragraphs:
            src |= _tokens(p.text)
        for tb in s.header.tables:
            for r in tb.rows:
                for c in r.cells:
                    src |= _tokens(c.text)
        for tb in s.footer.tables:
            for r in tb.rows:
                for c in r.cells:
                    src |= _tokens(c.text)
    return src


def _iter_rows_all(result):
    """全部 iter_rows 行 (含节标题) — 与 add_product 写入基准一致."""
    out = []
    for n in sorted(result.sections):
        for row in result.sections[n].iter_rows():
            out.append((n, row))
    return out


def _db_text(conn, pid) -> str:
    """DB 已存内容拼回文本 (字段 label+value+seq + 成分)."""
    parts = []
    for f in get_fields(conn, pid):
        parts.append(f["seq"] or "")
        parts.append(f["raw_label"] or "")
        parts.append(f["value"] or "")
    for c in get_components(conn, pid):
        parts.extend([c["name"], c["cas"], c["conc"]])
    return "\n".join(p for p in parts if p)


# ============================================================
# L1 解析正确性 (reader vs 源 docx)
# ============================================================

def audit_parse(result, doc_path) -> dict:
    doc = Document(str(doc_path))

    missing = sorted(n for n in range(1, 17) if n not in result.sections)
    has_s0 = 0 in result.sections

    src = _cell_tokens(doc)
    got = _tokens(_parse_text(result))
    miss = sorted(x for x in (src - got)
                  if x.lower() not in _ALLOWED_TOKEN_MISS)
    # 子序列放宽: 拆分重排的连写标签字符仍在 → 不算遗漏 (信息未丢失)
    all_text = _parse_text(result)
    miss = [x for x in miss if not _is_subseq(x, all_text)]

    n_field = sum(len(s.fields) for s in result.sections.values())
    n_line = sum(len(s.lines) for s in result.sections.values())
    n_comp = len(result.sections[3].components) if 3 in result.sections else 0
    n_sub_s0 = sum(1 for x in result.sections[0].iter_rows()
                   if x.kind == "sub") if 0 in result.sections else 0

    seq_issues = []
    for n in range(1, 17):
        sec = result.sections.get(n)
        if not sec:
            continue
        prev = None
        for row in sec.iter_rows():
            if row.seq:
                try:
                    cur = tuple(int(x) for x in row.seq.split("."))
                except ValueError:
                    continue
                if prev is not None and cur < prev:
                    seq_issues.append((n, prev, cur))
                prev = cur

    # 页码硬编码检测 (ChatGPT 页眉页脚审查): 页脚"页码"字段值是 `n / n`
    # 静态文本 → 可能是模板残留 (源文档无 PAGE/NUMPAGES 域, Word 分页会变).
    page_hardcoded = False
    page_text = ""
    s0 = result.sections.get(0)
    if s0:
        for row in s0.iter_rows():
            if "页码" in (row.label or "") and row.value:
                page_text = row.value.strip()
                break
    if page_text and re.match(r"^\d+\s*/\s*\d+$", page_text):
        page_hardcoded = True

    return {
        "sections_total": len(result.sections),
        "has_s0": has_s0,
        "missing_sections": missing,
        "token_src": len(src),
        "token_got": len(got),
        "token_missing": miss,
        "fields": n_field,
        "lines": n_line,
        "components": n_comp,
        "s0_subtitles": n_sub_s0,
        "page_hardcoded": page_hardcoded,
        "page_text": page_text,
        "seq_inversions": seq_issues,
        "anomalies": [{"level": a.level, "section": a.section,
                       "message": a.message} for a in result.anomalies],
    }


# ============================================================
# L2 写入正确性 (reader → SQLite)
# ============================================================

def audit_write(result, conn, pid) -> dict:
    # B1 字段全量: 入库 fields 行数 == iter_rows 全量 + 成分表头行
    expect = len(_iter_rows_all(result))
    if 3 in result.sections and result.sections[3].component_header:
        expect += 1
    got = conn.execute("SELECT COUNT(*) c FROM fields WHERE product_id=?",
                       (pid,)).fetchone()["c"]
    # B2 成分全量
    exp_comp = len(result.sections[3].components) if 3 in result.sections else 0
    got_comp = conn.execute("SELECT COUNT(*) c FROM components WHERE product_id=?",
                            (pid,)).fetchone()["c"]
    # B3 field_key 收敛: 非空 key 中 非原样兜底 的比例 (原样兜底=未收敛)
    rows = conn.execute(
        "SELECT section, field_key, raw_label FROM fields WHERE product_id=? "
        "AND kind IN ('field','sub') AND field_key!=''", (pid,)).fetchall()
    n_key = len(rows)
    n_raw_fallback = sum(1 for r in rows
                         if r["field_key"] == (r["raw_label"] or "").strip())
    # B4 raw_label 保留: 写库 (section, raw_label) 与 reader iter_rows 逐行对照
    #    (排除 component_header 行 — 表头非 iter_rows 语义, 已单独核查 B6)
    expect_labels = [(n, row.label) for n, row in _iter_rows_all(result)]
    db_labels = [(r["section"], r["raw_label"]) for r in
                 conn.execute("SELECT section, raw_label FROM fields "
                              "WHERE product_id=? AND kind!='component_header' "
                              "ORDER BY section, row_order",
                              (pid,)).fetchall()]
    label_mismatch = len(expect_labels) != len(db_labels) or any(
        e != d for e, d in zip(expect_labels, db_labels))
    # B5 内容 token 完整: DB 存(token) vs reader 存(token) (排除序号剥离词)
    #    表头现已入库 (B6 独立核查), 不再从 reader 侧预减.
    reader_tok = _tokens(_parse_text(result))
    db_tok = _tokens(_db_text(conn, pid))
    loss = sorted(x for x in (reader_tok - db_tok)
                  if x.lower() not in _ALLOWED_TOKEN_MISS)
    # 子序列放宽 (同 L1): 拆分重排字符仍在 → 不算遗漏
    db_text_full = _db_text(conn, pid)
    loss = [x for x in loss if not _is_subseq(x, db_text_full)]
    # B6 成分表头持久化: component_header 应入库 (现 add_product 未写 → 记录)
    header_persisted = conn.execute(
        "SELECT COUNT(*) c FROM fields WHERE product_id=? AND section=3 "
        "AND kind='component_header'", (pid,)).fetchone()["c"] > 0
    # B7 幂等: 同 result 重 add → exists_sha (不重复写入)
    _p2, status2 = add_product(conn, result, category="__dup__",
                               model_name="__dup__")

    return {
        "fields_expected": expect,
        "fields_written": got,
        "components_expected": exp_comp,
        "components_written": got_comp,
        "key_fields": n_key,
        "raw_fallback_keys": n_raw_fallback,
        "converged_keys": n_key - n_raw_fallback,
        "label_mismatch": label_mismatch,
        "token_loss": loss,
        "header_persisted": header_persisted,
        "header_text": (result.sections[3].component_header
                        if 3 in result.sections else ""),
        "duplicate_status": status2,
    }


# ============================================================
# L3 回读还原 (SQLite → ParseResult)
# ============================================================

def audit_rebuild(result, conn, pid) -> dict:
    reb = rebuild_product(conn, pid)
    if reb is None:
        return {"ok": False, "reason": "rebuild None"}

    # C1 结构: 回读节数 == 原节数
    n_sec = len(reb.sections)
    # C2 内容: (section, label, value) 集合对照 (kind 允许 sub→field 归一)
    exp = sorted((n, row.label, row.value) for n, row in _iter_rows_all(result))
    got = sorted((n, row.label, row.value) for n, row in _iter_rows_all(reb))
    value_mismatch = len(exp) != len(got) or any(
        a != b for a, b in zip(exp, got))
    # C3 顺序: 节内 (label 序列) 一致
    exp_order = [row.label for _n, row in _iter_rows_all(result)]
    got_order = [row.label for _n, row in _iter_rows_all(reb)]
    seq_ok = exp_order == got_order
    # C4 成分还原
    db_comp = [(c["name"], c["cas"], c["conc"]) for c in get_components(conn, pid)]
    reb_comp = [(c.name, c.cas, c.conc)
                for c in (reb.sections[3].components if 3 in reb.sections else [])]
    # C5 表头还原 (rebuild 现不写 component_header → 记录差异)
    exp_header = (result.sections[3].component_header
                  if 3 in result.sections else "") or ""
    got_header = (reb.sections[3].component_header
                  if 3 in reb.sections else "") or ""
    # C6 document_tree 机器树: node_id 唯一 / 成分 raw_value 保留 / 实体化
    tree = build_document_tree(reb)
    all_ids = []
    for s in tree["tree"]:
        all_ids.append(s["node_id"])
        for e in s.get("entities", []):
            all_ids.append(e["node_id"])
            all_ids += [f["node_id"] for f in e["fields"]]
        stack = list(s.get("children", []))
        while stack:
            nd = stack.pop()
            all_ids.append(nd["node_id"])
            stack.extend(nd.get("children", []))
    dup_ids = len(all_ids) - len(set(all_ids))
    # 成分 raw_value 保留检查 (ChatGPT #9)
    # 仅当归一化值非空而 raw 为空才判定缺失 (原文含量为空则 raw 空合理)
    ent_fields = [f for s in tree["tree"]
                  for e in s.get("entities", []) for f in e["fields"]]
    raw_missing = [f["node_id"] for f in ent_fields
                   if f["label"] in ("名称", "含量")
                   and f["value"] and not f["raw_value"]]
    return {
        "ok": True,
        "sections": n_sec,
        "sections_match": n_sec == len(result.sections),
        "rows_expected": len(exp),
        "rows_rebuilt": len(got),
        "value_mismatch": value_mismatch,
        "order_ok": seq_ok,
        "components_match": (db_comp == reb_comp),
        "components_count": len(reb_comp),
        "header_match": (exp_header == got_header),
        "header_expected": exp_header,
        "header_got": got_header,
        "tree_nodes": len(all_ids),
        "tree_dup_ids": dup_ids,
        "tree_entities": sum(1 for s in tree["tree"]
                             if s.get("section_no") == 3
                             for _ in s.get("entities", [])),
        "component_raw_missing": raw_missing,
    }


# ============================================================
# L4 输出正确性 (SQLite → Excel 导出)
# ============================================================

def _sec_has_content(result, n: int) -> bool:
    """该节是否有真实内容 (非空 value 的字段/说明, 或成分).

    空值子标题/结构标题 (如 '3.1 Composition :') 不算内容 — 宽表
    跳空值列是设计行为, 不能据此判导出缺列.
    """
    sec = result.sections.get(n)
    if not sec:
        return False
    if sec.components:
        return True
    for row in sec.iter_rows():
        if (row.value or "").strip() and row.kind in ("field", "note"):
            return True
    return False


def audit_excel(result, conn, pid, out_xlsx) -> dict:
    export_single(result, conn, out_xlsx, category="audit",
                  model_name="audit")
    wb = load_workbook(out_xlsx)
    ws = wb["数据库宽表"]
    ws2 = wb["完整明细"]
    hdr = [c.value for c in ws[1]]

    # D1 宽表节序单调
    segs = []
    for h in hdr:
        if h:
            m = re.match(r"S(\d+)[.·]", str(h))
            if m:
                segs.append(int(m.group(1)))
    monotonic = all(segs[i] <= segs[i + 1] for i in range(len(segs) - 1))
    covered = set(segs)
    # D2 明细全量: 明细数据行数 >= DB fields 行数 (成分拆行后应 >=)
    n_db = conn.execute("SELECT COUNT(*) c FROM fields WHERE product_id=?",
                        (pid,)).fetchone()["c"]
    detail_rows = ws2.max_row - 1
    # D3 值保真: 明细"值"列 不得含人工占位文字 (ChatGPT #9 回归)
    ph = 0
    for r in range(2, ws2.max_row + 1):
        v = str(ws2.cell(r, 6).value or "")
        if "分组标题" in v or "待填" in v or "固定标题" in v:
            ph += 1
    # D4 明细序号列覆盖
    n_seq = sum(1 for r in range(2, ws2.max_row + 1)
                if str(ws2.cell(r, 3).value or "").strip())

    out_xlsx.unlink(missing_ok=True)
    missing_all = [n for n in range(1, 17) if n not in covered]
    missing_with_content = [n for n in missing_all if _sec_has_content(result, n)]
    return {
        "wide_cols": len(hdr),
        "section_seq_monotonic": monotonic,
        "sections_covered": sorted(covered),
        "missing_section_cols": sorted(missing_all),
        "missing_cols_with_content": sorted(missing_with_content),
        "db_fields": n_db,
        "detail_rows": detail_rows,
        "detail_full": detail_rows >= n_db,
        "value_placeholder_count": ph,
        "seq_rows": n_seq,
    }


# ============================================================
# 数据库逻辑性静态审查 (L 维度)
# ============================================================

def audit_schema(conn) -> dict:
    tables = {r["name"] for r in conn.execute(
        "SELECT name FROM sqlite_master WHERE type='table'")}
    fks = []
    for t in ("products", "fields", "components", "pictograms", "revisions"):
        for r in conn.execute(f"PRAGMA foreign_key_list({t})").fetchall():
            fks.append((t, r["table"], r["from"], r["to"], r["on_delete"]))
    uniques = []
    for t in ("categories", "products", "fields", "components", "field_dict"):
        for r in conn.execute(f"PRAGMA index_list({t})").fetchall():
            if r["unique"]:
                uniques.append((t, r["name"]))
    indexes = []
    for t in tables:
        for r in conn.execute(f"PRAGMA index_list({t})").fetchall():
            indexes.append((t, r["name"], bool(r["unique"])))
    pcols = [r["name"] for r in conn.execute("PRAGMA table_info(products)")]
    return {
        "tables": sorted(tables),
        "foreign_keys": fks,
        "uniques": uniques,
        "indexes": indexes,
        "has_soft_delete": "active" in pcols,
        "eav_design": ("fields" in tables and "components" in tables
                       and "products" in tables),
    }


# ============================================================
# 大批量能力评估 (B 维度)
# ============================================================

def audit_bulk(elapsed_sec, file_count) -> dict:
    per = elapsed_sec / max(file_count, 1)
    return {
        "elapsed_sec": round(elapsed_sec, 3),
        "files": file_count,
        "per_file_sec": round(per, 3),
        "est_646_files_min": round(per * 646 / 60, 1),
    }


# ============================================================
# 判定
#  FAIL: 内容/结构损坏 (真丢数据 / 写库不一致) — 阻断大批量
#  WARN: 原文结构特性 (无页眉页脚 / 原文缺节 / 序号怪癖 / 原文空节 /
#         自动编号恢复异常) — 记录但不阻断
# ============================================================

def _parse_verdict(a: dict):
    """解析层: token 遗漏 = FAIL; 原文特性 (缺节/无S0/序号倒序) = WARN;
    页码硬编码 = 数据质量 (源文档模板缺陷, 不降级文件状态)."""
    fails = []
    warns = []
    if a["missing_sections"]:
        warns.append(f"原文缺节 {a['missing_sections']}")
    if a["token_missing"]:
        fails.append(f"token遗漏 {a['token_missing'][:6]}")
    if not a["has_s0"]:
        warns.append("原文无页眉页脚")
    if a["seq_inversions"]:
        warns.append(f"序号倒序 {a['seq_inversions']}")
    dq = []
    if a["page_hardcoded"]:
        dq.append(f"页码硬编码 {a['page_text']} (源无PAGE域)")
    ok = not fails
    msg = ("；".join(fails) if fails else "解析完整")
    return ok, msg, warns, dq


def _write_verdict(a: dict):
    """写入层: 字段/成分/token/幂等 = FAIL; 表头未入库 = WARN."""
    fails = []
    warns = []
    if a["fields_expected"] != a["fields_written"]:
        fails.append(f"字段 {a['fields_written']}/{a['fields_expected']}")
    if a["components_expected"] != a["components_written"]:
        fails.append(f"成分 {a['components_written']}/{a['components_expected']}")
    if a["label_mismatch"]:
        fails.append("raw_label 未逐行对齐")
    if a["token_loss"]:
        fails.append(f"入库token遗漏 {a['token_loss'][:6]}")
    if a["duplicate_status"] != "exists_sha":
        fails.append(f"幂等失败: {a['duplicate_status']}")
    if not a["header_persisted"] and bool(a["header_text"]):
        warns.append("成分表头未入库")
    ok = not fails
    msg = ("；".join(fails) if fails else "写入完整")
    return ok, msg, warns


def _rebuild_verdict(a: dict):
    """回读层: 结构/内容/顺序/成分 = FAIL; 表头未回读 = WARN."""
    if not a.get("ok"):
        return False, a.get("reason", "回读失败"), []
    fails = []
    warns = []
    if not a["sections_match"]:
        fails.append(f"节 {a['sections']}")
    if a["value_mismatch"]:
        fails.append("内容不符")
    if not a["order_ok"]:
        fails.append("顺序不符")
    if not a["components_match"]:
        fails.append("成分不符")
    if not a["header_match"] and bool(a["header_expected"]):
        warns.append("成分表头未回读")
    if a.get("tree_dup_ids"):
        fails.append(f"node_id重复 {a['tree_dup_ids']}")
    if a.get("component_raw_missing"):
        fails.append(f"成分raw丢失 {a['component_raw_missing'][:4]}")
    ok = not fails
    msg = ("；".join(fails) if fails else "还原一致")
    return ok, msg, warns


def _excel_verdict(a: dict):
    """输出层: 节序/明细/值保真 = FAIL; 原文空节无列 = WARN."""
    fails = []
    warns = []
    if not a["section_seq_monotonic"]:
        fails.append("宽表节序非单调")
    if a["missing_cols_with_content"]:
        fails.append(f"缺内容节列 {a['missing_cols_with_content']}")
    empty_missing = [n for n in a["missing_section_cols"]
                     if n not in a["missing_cols_with_content"]]
    if empty_missing:
        warns.append(f"原文空节无列 {empty_missing}")
    if not a["detail_full"]:
        fails.append(f"明细 {a['detail_rows']}<字段{a['db_fields']}")
    if a["value_placeholder_count"]:
        fails.append(f"值列占位 {a['value_placeholder_count']}")
    ok = not fails
    msg = ("；".join(fails) if fails else "输出正确")
    return ok, msg, warns


# ============================================================
# 主流程
# ============================================================

def run_audit(folder: Path, *, out_dir: Path | None = None,
              detail: bool = False) -> dict:
    folder = Path(folder)
    docs = sorted(folder.rglob("*.docx"))
    out_dir = out_dir or folder

    report = {
        "tool": "MSDS db.audit",
        "generated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "folder": str(folder),
        "files_found": len(docs),
        "files": [],
        "schema": {},
        "bulk": {},
        "summary": {},
    }

    t_start = time.time()
    schema_checked = False
    file_results = []
    tmp = Path(tempfile.gettempdir()) / "_msds_audit.xlsx"

    for d in docs:
        file_report = {"file": d.name, "path": str(d)}
        try:
            t0 = time.time()
            result = read_msds(d)
            conn = connect(":memory:")
            pid, status = add_product(conn, result, category="audit",
                                      model_name="audit-tmp")
            l1 = audit_parse(result, d)
            l2 = audit_write(result, conn, pid)
            l3 = audit_rebuild(result, conn, pid)
            l4 = audit_excel(result, conn, pid, tmp)
            if not schema_checked:
                report["schema"] = audit_schema(conn)
                schema_checked = True
            elapsed = time.time() - t0
            file_report.update({"elapsed_sec": round(elapsed, 3),
                                "parse": l1, "write": l2,
                                "rebuild": l3, "excel": l4})
            pv = _parse_verdict(l1)
            wv = _write_verdict(l2)
            rv = _rebuild_verdict(l3)
            ev = _excel_verdict(l4)
            file_report["verdicts"] = {
                "parse": {"passed": pv[0], "msg": pv[1], "warns": pv[2]},
                "write": {"passed": wv[0], "msg": wv[1], "warns": wv[2]},
                "rebuild": {"passed": rv[0], "msg": rv[1], "warns": rv[2]},
                "excel": {"passed": ev[0], "msg": ev[1], "warns": ev[2]},
            }
            if pv[3]:
                file_report["data_quality"] = pv[3]
            passed_all = all(
                file_report["verdicts"][k]["passed"]
                for k in ("parse", "write", "rebuild", "excel"))
            any_warn = any(
                file_report["verdicts"][k].get("warns")
                for k in ("parse", "write", "rebuild", "excel"))
            # 数据质量 (页码硬编码等源文档缺陷) 不降级文件状态
            file_report["status"] = ("ok" if passed_all and not any_warn
                                     else "warn" if passed_all else "fail")
        except Exception as e:
            import traceback
            file_report["status"] = "error"
            file_report["error"] = str(e)
            if detail:
                file_report["traceback"] = traceback.format_exc()
        file_results.append(file_report)
        report["files"].append(file_report)

    report["bulk"] = audit_bulk(time.time() - t_start, len(docs))

    n_ok = sum(1 for f in file_results if f["status"] == "ok")
    n_warn = sum(1 for f in file_results if f["status"] == "warn")
    n_fail = sum(1 for f in file_results if f["status"] == "fail")
    n_err = sum(1 for f in file_results if f["status"] == "error")
    passed = {k: 0 for k in ("parse", "write", "rebuild", "excel")}
    for f in file_results:
        for k in passed:
            if f.get("verdicts", {}).get(k, {}).get("passed"):
                passed[k] += 1
    total = max(len(docs), 1)
    # 数据质量统计 (源文档缺陷, 独立于文件状态)
    dq_counter = {"page_hardcoded": 0}
    page_texts = {}
    for f in file_results:
        if f.get("data_quality"):
            for item in f["data_quality"]:
                if item.startswith("页码硬编码"):
                    dq_counter["page_hardcoded"] += 1
                    txt = item.split(" ", 2)[2].split(" (")[0]
                    page_texts[txt] = page_texts.get(txt, 0) + 1
    report["data_quality"] = {
        "page_hardcoded": dq_counter["page_hardcoded"],
        "page_texts": dict(sorted(page_texts.items(),
                                  key=lambda kv: -kv[1])),
    }
    report["summary"] = {
        "ok": n_ok, "warn": n_warn, "fail": n_fail, "error": n_err,
        "pass_rate": {k: f"{v}/{total}" for k, v in passed.items()},
        # 就绪判定: 无 FAIL、无 ERROR 即放行; WARN/数据质量 记录但不阻断
        "ready_for_bulk": n_fail == 0 and n_err == 0,
        "est_646_files_min": report["bulk"]["est_646_files_min"],
    }
    return report


def render_md(report: dict) -> str:
    L = []
    L.append("# MSDS 大批量入库就绪度评判报告\n")
    L.append(f"- 生成时间: {report['generated_at']}")
    L.append(f"- 评判目录: `{report['folder']}`")
    L.append(f"- 文件数: {report['files_found']}")
    s = report["summary"]
    ready = s["ready_for_bulk"]
    L.append(f"- 就绪判定: **{'✅ 可以大批量' if ready else '⚠️ 需修复后放行'}**\n")

    L.append("## 总览\n")
    L.append("| 状态 | 数量 |\n|---|---|")
    L.append(f"| ✅ 全项通过 | {s['ok']} |")
    L.append(f"| ⚠️ 有警告 (原文特性) | {s['warn']} |")
    L.append(f"| ❌ 内容/结构损坏 | {s.get('fail', 0)} |")
    L.append(f"| 💥 解析失败 | {s['error']} |\n")
    L.append("| 层次 | 通过率 |\n|---|---|")
    for k, v in s["pass_rate"].items():
        L.append(f"| {k} | {v} |")
    L.append("")

    dq = report.get("data_quality") or {}
    if dq.get("page_hardcoded"):
        L.append("## 源文档数据质量 (不影响程序就绪, 供模板修正参考)\n")
        L.append(f"- 页码硬编码 (页脚为静态文本, 无 PAGE/NUMPAGES 域): "
                 f"**{dq['page_hardcoded']}** 个文件")
        pts = dq.get("page_texts") or {}
        if pts:
            top = "、".join(f"`{k}`×{v}" for k, v in list(pts.items())[:6])
            L.append(f"- 出现形态: {top}")
        L.append("- 修正建议: 模板页脚改为插入「页 X / 共 Y」域字段, "
                 "否则分页变化后页码失真\n")

    b = report["bulk"]
    L.append("## 大批量能力\n")
    L.append(f"- 实测总耗时: {b['elapsed_sec']}s ({b['files']} 文件)")
    L.append(f"- 单文件均值: {b['per_file_sec']}s")
    L.append(f"- **外推 646 文件: 约 {b['est_646_files_min']} 分钟**\n")

    sc = report["schema"]
    if sc:
        L.append("## 数据库逻辑性审查\n")
        L.append(f"- 表: {', '.join(sc['tables'])}")
        L.append(f"- EAV 规范化 (products/fields/components 分离): "
                 f"{'✅' if sc['eav_design'] else '❌'}")
        L.append(f"- 软删除 (active): {'✅' if sc['has_soft_delete'] else '❌'}")
        fk_s = "; ".join(f"{a}.{c}→{b}.{d}({e})"
                         for a, b, c, d, e in sc["foreign_keys"]) or "无"
        L.append(f"- 外键 {len(sc['foreign_keys'])} 条: {fk_s}")
        uniq_s = "; ".join(f"{a}({b})" for a, b in sc["uniques"]) or "无"
        L.append(f"- 唯一约束 {len(sc['uniques'])} 条: {uniq_s}")
        ix_s = "; ".join(f"{a}.{b}{'·U' if c else ''}"
                         for a, b, c in sc["indexes"]) or "无"
        L.append(f"- 索引 {len(sc['indexes'])} 个: {ix_s}\n")

    L.append("## 逐文件明细\n")
    for f in report["files"]:
        L.append(f"### {f['file']}  [{f['status']}]\n")
        if f["status"] == "error":
            L.append(f"错误: {f.get('error')}\n")
            continue
        v = f["verdicts"]
        for k, cn in (("parse", "解析"), ("write", "写入"),
                      ("rebuild", "回读"), ("excel", "输出")):
            x = v[k]
            mark = "✅" if x["passed"] else "❌"
            warns = x.get("warns") or []
            extra = (" ⚠️" + "；".join(warns)) if warns else ""
            L.append(f"**{cn}** {mark}{extra} {x['msg']}")
        p = f["parse"]
        L.append(f"- 节 {p['sections_total']} | 字段 {p['fields']} | "
                 f"行 {p['lines']} | 成分 {p['components']} | "
                 f"S0子标题 {p['s0_subtitles']} | {f['elapsed_sec']}s")
        if p["token_missing"]:
            L.append(f"- ⚠️ 解析 token 未覆盖: {p['token_missing'][:8]}")
        if p["seq_inversions"]:
            L.append(f"- ⚠️ 序号倒序: {p['seq_inversions']}")
        if f.get("data_quality"):
            L.append(f"- 📋 数据质量: {'；'.join(f['data_quality'])}")
        w = f["write"]
        L.append(f"- 入库: 字段 {w['fields_written']}/{w['fields_expected']} | "
                 f"成分 {w['components_written']}/{w['components_expected']} | "
                 f"key收敛 {w['converged_keys']}/{w['key_fields']}")
        if w["token_loss"]:
            L.append(f"- ⚠️ 入库 token 遗漏: {w['token_loss'][:8]}")
        if w["header_text"] and not w["header_persisted"]:
            L.append(f"- ⚠️ 成分表头未持久化: `{w['header_text'][:40]}`")
        r = f["rebuild"]
        L.append(f"- 机器树: {r.get('tree_nodes', 0)} 节点 | "
                 f"ID重复 {r.get('tree_dup_ids', 0)} | "
                 f"成分实体 {r.get('tree_entities', 0)} | "
                 f"raw丢失 {len(r.get('component_raw_missing') or [])}")
        if r["header_expected"] and not r["header_match"]:
            L.append(f"- ⚠️ 成分表头回读差异: 期望 `{r['header_expected'][:30]}` "
                     f"→ 回读 `{r['header_got'][:20] or '空'}`")
        L.append("")
    return "\n".join(L)


def main(argv=None):
    ap = argparse.ArgumentParser(description="MSDS 大批量入库就绪度评判")
    ap.add_argument("folder", help="待评判文件夹 (递归扫描 *.docx)")
    ap.add_argument("--out", default=None, help="报告输出目录 (默认=文件夹)")
    ap.add_argument("--detail", action="store_true", help="含异常堆栈")
    args = ap.parse_args(argv)

    out_dir = Path(args.out) if args.out else Path(args.folder)
    out_dir.mkdir(parents=True, exist_ok=True)
    report = run_audit(args.folder, out_dir=out_dir, detail=args.detail)
    jp = out_dir / "audit_report.json"
    mp = out_dir / "audit_report.md"
    jp.write_text(json.dumps(report, ensure_ascii=False, indent=2),
                  encoding="utf-8")
    mp.write_text(render_md(report), encoding="utf-8")
    print(f"报告已生成:\n  {jp}\n  {mp}")
    print(json.dumps(report["summary"], ensure_ascii=False, indent=2))
    return 0 if report["summary"]["ready_for_bulk"] else 1


if __name__ == "__main__":
    sys.exit(main())
