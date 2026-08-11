# -*- coding: utf-8 -*-
"""覆写基座: 在保留原文格式的前提下, 对 Word 文档做最小化覆写.

覆写原则 (用户需求):
  - 只允许 添加 / 删减 / 编辑 部分内容, 并覆写进文档
  - 不改变原文格式 (覆写内容继承原文格式)
  - 基于内化默认模板 (templates/MSDS_CN 国彩 模板.docx, 字节级一致)

技术要点:
  - 替换单元格文本时**逐 run 覆写** (保留首个 run 的字体/字号/粗细),
    绝不使用 cell.text = ... (会清空段落与 run 格式, 丢失格式).
  - 多段落单元格 (含 \n) 保持段落数, 逐段写入, 不新增/删除段落.
  - 整表行/列的增删 (添加/删减) 走 python-docx 的 table 行操作,
    复用临近行的格式.
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.table import Table, _Row
from docx.text.paragraph import Paragraph

from .docx_reader import (
    TEMPLATE_PATH,
    _dedupe_row,
    is_component_header_row,
    is_section_title,
)


# ---------------- 单元格格式保留覆写 ----------------

def _cell_paragraphs(cell) -> list[Paragraph]:
    """取单元格段落 (兼容空段落)."""
    return list(cell.paragraphs)


def set_cell_text(cell, text: str) -> None:
    """把单元格文本覆写为 text, 保留段落结构与首个 run 的格式.

    - text 含 \n → 保持多段落, 逐段写入
    - 每段保留首个 run 的字体/字号/粗细; 无 run 时新建 run 继承段落样式
    - 段数不足补段落; 段数多余只写前几段 (不删除, 保守)
    """
    lines = (text or "").split("\n")
    paras = _cell_paragraphs(cell)
    for i, line in enumerate(lines):
        para = paras[i] if i < len(paras) else cell.add_paragraph()
        _set_paragraph_text(para, line)


def _set_paragraph_text(para: Paragraph, text: str) -> None:
    """覆写段落文本, 保留首个 run 格式 (若存在)."""
    # 取首个 run 作格式模板
    fmt = None
    if para.runs:
        fmt = para.runs[0]
    # 清空段落现有 runs
    for r in list(para.runs):
        r._r.getparent().remove(r._r)
    if not text:
        return
    if fmt is not None:
        # 复用首个 run (保留其格式), 改写其文本
        run = para.add_run(text)
        _copy_run_format(fmt, run)
    else:
        para.add_run(text)


def _copy_run_format(src, dst) -> None:
    """把源 run 的字体属性复制到目标 run (无格式时才补默认)."""
    if src.font.name:
        dst.font.name = src.font.name
    if src.font.size:
        dst.font.size = src.font.size
    if src.font.bold is not None:
        dst.font.bold = src.font.bold
    if src.font.italic is not None:
        dst.font.italic = src.font.italic
    if src.font.underline is not None:
        dst.font.underline = src.font.underline
    if src.font.color and src.font.color.rgb:
        dst.font.color.rgb = src.font.color.rgb


# ---------------- 整表行增删 (添加/删减) ----------------

def _find_table(doc: Document, section_num: int) -> tuple[Table, int] | None:
    """按节号定位表格. 返回 (table, 表格内节标题行索引) 或 None."""
    for tb in doc.tables:
        for ri, row in enumerate(tb.rows):
            cells, _ = _dedupe_row(row)
            if cells and cells[0]:
                n, _ = is_section_title(cells[0])
                if n == section_num:
                    return tb, ri
    return None


def add_table_row(table: Table, index: int | None = None, template_row: int = 0) -> _Row:
    """在表格中新增一行, 复制 template_row 的格式 (继承原文格式)."""
    import copy
    # python-docx: 复制 XML 行元素
    src_tr = table.rows[template_row]._tr
    new_tr = copy.deepcopy(src_tr)
    if index is None:
        src_tr.addnext(new_tr)
    else:
        table.rows[index]._tr.addprevious(new_tr)
    return _Row(new_tr, table)


def delete_table_row(table: Table, index: int) -> None:
    """删除表格第 index 行."""
    tr = table.rows[index]._tr
    tr.getparent().remove(tr)


# ---------------- 顶层覆写入口 ----------------

def overwrite_doc(src: str | Path, changes: dict[tuple[int, int, int], str],
                  out: str | Path, component_index: bool = False) -> None:
    """把 src 文档按 changes 覆写并另存为 out.

    changes: {(节号, 行位置, 单元格列号): 新文本}
      - component_index=False (默认): 行位置 = 该节表格内的原始行号 (0 起,
        含节标题行). 适合字段行覆写 (如 S1 中文名称行=2).
      - component_index=True: 行位置 = 该节**成分表的第 N 个数据行** (0 起),
        自动跳过表头/产品类型/列标题行. 适合 S3 成分表覆写 (不随表格行号偏移).
      单元格列号: 0 起 (S3 成分表 0=名称 1=CAS 2=含量).
    其余内容原样保留 → 格式零丢失.
    """
    doc = Document(str(src))
    # 预解析源文档, 判定哪些节是成分表 (成分索引仅对这些节生效)
    comp_sections: set[int] = set()
    if component_index:
        from .docx_reader import read_msds
        for n, s in read_msds(src).sections.items():
            if s.is_component_table:
                comp_sections.add(n)

    for (sec_num, row_pos, col_idx), new_text in changes.items():
        tb, sec_row = _find_table(doc, sec_num)
        if tb is None:
            raise ValueError(f"未找到第{sec_num}节表格")
        if component_index and sec_num in comp_sections:
            row_idx = _component_data_row(tb, sec_row, row_pos)
        else:
            row_idx = row_pos
        cell = tb.rows[row_idx].cells[col_idx]
        set_cell_text(cell, new_text)
    doc.save(str(out))


def _component_data_row(table: Table, sec_row: int, comp_idx: int) -> int:
    """返回成分表第 comp_idx 个数据行在 table 中的实际行号.

    跳过: 节标题行 / 产品类型行 / 成分列标题行 (['成分',''] / ['成分 /','']) /
          成分表头行 (化学品名称|CAS|含量) / 单列说明行.
    """
    data_rows: list[int] = []
    for ri in range(sec_row + 1, len(table.rows)):
        cells, _ = _dedupe_row(table.rows[ri])
        if not cells or not cells[0]:
            continue
        first = cells[0].strip()
        # 跳过标题类行
        if first.startswith("产品类型"):
            continue
        if first.rstrip(" /／ ") in ("成分", "组分", "成分/组成", "危险成分", "名称", "化学品名称"):
            continue
        if is_component_header_row(cells):
            continue
        if len(cells) < 3:
            continue  # 单列说明行
        data_rows.append(ri)
    if comp_idx >= len(data_rows):
        raise ValueError(f"成分索引 {comp_idx} 超出 (仅 {len(data_rows)} 个成分数据行)")
    return data_rows[comp_idx]


def copy_template(out: str | Path) -> None:
    """把内化默认模板二进制复制为 out (格式零丢失, 供覆写前做副本)."""
    shutil.copy2(TEMPLATE_PATH, str(out))
